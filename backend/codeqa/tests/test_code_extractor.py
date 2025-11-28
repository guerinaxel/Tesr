from __future__ import annotations

from pathlib import Path
from tempfile import TemporaryDirectory
from unittest.mock import Mock, patch

from django.test import SimpleTestCase
from docx import Document
from PyPDF2 import PdfWriter

from codeqa.code_extractor import collect_code_chunks, extract_pdf_text


PDF_CONTENT = b"""%PDF-1.4
1 0 obj
<< /Type /Catalog /Pages 2 0 R >>
endobj
2 0 obj
<< /Type /Pages /Kids [3 0 R] /Count 1 >>
endobj
3 0 obj
<< /Type /Page /Parent 2 0 R /MediaBox [0 0 300 144] /Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>
endobj
4 0 obj
<< /Length 44 >>
stream
BT
/F1 24 Tf
100 100 Td
(Test PDF) Tj
ET
endstream
endobj
5 0 obj
<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>
endobj
xref
0 6
0000000000 65535 f 
0000000010 00000 n 
0000000079 00000 n 
0000000178 00000 n 
0000000409 00000 n 
0000000502 00000 n 
trailer
<< /Root 1 0 R /Size 6 >>
startxref
582
%%EOF
"""


class CollectCodeChunksTests(SimpleTestCase):
    def setUp(self) -> None:
        self.tmpdir = TemporaryDirectory()
        self.root_path = Path(self.tmpdir.name)
        self.addCleanup(self.tmpdir.cleanup)

    def test_collects_pdf_content(self) -> None:
        # Arrange
        pdf_path = self.root_path / "sample.pdf"
        pdf_path.write_bytes(PDF_CONTENT)

        # Act
        chunks = collect_code_chunks(self.root_path)

        # Assert
        self.assertTrue(
            any("File: sample.pdf" in chunk and "Test PDF" in chunk for chunk in chunks),
        )

    def test_collects_docx_content(self) -> None:
        # Arrange
        docx_path = self.root_path / "sample.docx"
        document = Document()
        document.add_paragraph("Hello Word Doc")
        document.save(docx_path)

        # Act
        chunks = collect_code_chunks(self.root_path)

        # Assert
        self.assertTrue(
            any("File: sample.docx" in chunk and "Hello Word Doc" in chunk for chunk in chunks),
        )

    def test_python_ast_chunking_with_nested_defs(self) -> None:
        # Arrange
        python_path = self.root_path / "sample.py"
        python_path.write_text(
            """
class Outer:
    def method_one(self):
        def inner():
            return "hi"
        return inner()


def standalone(value):
    return value * 2
""".strip()
        )

        # Act
        chunks = collect_code_chunks(self.root_path)

        # Assert
        self.assertTrue(any("class Outer" in chunk for chunk in chunks))
        self.assertTrue(any("def standalone" in chunk for chunk in chunks))

    def test_typescript_ast_chunking_and_large_nodes(self) -> None:
        # Arrange
        ts_path = self.root_path / "sample.ts"
        big_body = "\n".join("console.log('line')" for _ in range(200))
        ts_path.write_text(
            f"""
class Service {{
    method() {{
{big_body}
    }}
}}

function helper(x: number) {{
    return x + 1;
}}
""".strip()
        )

        # Act
        chunks = collect_code_chunks(self.root_path)

        # Assert
        ts_chunks = [chunk for chunk in chunks if "sample.ts" in chunk]
        self.assertTrue(any("class Service" in chunk for chunk in ts_chunks))
        self.assertTrue(any("function helper" in chunk for chunk in ts_chunks))
        self.assertTrue(any(len(chunk.splitlines()) > 50 for chunk in ts_chunks))

    def test_typescript_falls_back_to_line_chunking_on_error(self) -> None:
        # Arrange
        ts_path = self.root_path / "broken.ts"
        ts_path.write_text("function example() { return 'oops' " " ")

        with patch("codeqa.code_extractor.chunk_code_with_ast", side_effect=ValueError("parse err")):
            with patch(
                "codeqa.code_extractor.chunk_text",
                wraps=collect_code_chunks.__globals__["chunk_text"],
            ) as mock_chunk:
                # Act
                chunks = collect_code_chunks(self.root_path)

        # Assert
        ts_chunks = [chunk for chunk in chunks if "broken.ts" in chunk]
        self.assertTrue(ts_chunks)
        self.assertTrue(any("oops" in chunk for chunk in ts_chunks))
        self.assertGreaterEqual(mock_chunk.call_count, 1)

    def test_html_uses_text_chunking(self) -> None:
        # Arrange
        html_path = self.root_path / "index.html"
        repeated = "\n".join(f"<p>line {i}</p>" for i in range(100))
        html_path.write_text(f"<html><body>\n{repeated}\n</body></html>")

        # Act
        chunks = collect_code_chunks(self.root_path)

        # Assert
        html_chunks = [chunk for chunk in chunks if "index.html" in chunk]
        self.assertTrue(html_chunks)
        self.assertTrue(any("<html>" in chunk for chunk in html_chunks))
        self.assertGreater(len(html_chunks), 1)

    def test_pdf_and_docx_are_chunked_after_extraction(self) -> None:
        # Arrange
        pdf_path = self.root_path / "sample.pdf"
        pdf_path.write_bytes(PDF_CONTENT)
        docx_path = self.root_path / "sample.docx"
        document = Document()
        document.add_paragraph("Hello Word Doc")
        document.save(docx_path)

        with patch(
            "codeqa.code_extractor.chunk_text",
            wraps=collect_code_chunks.__globals__["chunk_text"],
        ) as mock_chunk:
            # Act
            chunks = collect_code_chunks(self.root_path)

        # Assert
        self.assertTrue(any("File: sample.pdf" in chunk for chunk in chunks))
        self.assertTrue(any("File: sample.docx" in chunk for chunk in chunks))
        self.assertGreaterEqual(mock_chunk.call_count, 2)


class ExtractPdfTextTests(SimpleTestCase):
    def setUp(self) -> None:
        self.tmpdir = TemporaryDirectory()
        self.root_path = Path(self.tmpdir.name)
        self.addCleanup(self.tmpdir.cleanup)

    def test_extracts_text_pdf_without_ocr(self) -> None:
        # Arrange
        pdf_path = self.root_path / "text.pdf"
        pdf_path.write_bytes(PDF_CONTENT)

        # Act
        text = extract_pdf_text(pdf_path)

        # Assert
        self.assertIn("Test PDF", text)

    def test_extracts_pdf_text_from_image_pages_via_ocr(self) -> None:
        # Arrange
        pdf_path = self.root_path / "image.pdf"
        writer = PdfWriter()
        writer.add_blank_page(width=200, height=200)
        with pdf_path.open("wb") as pdf_file:
            writer.write(pdf_file)

        mock_convert_from_path = Mock(return_value=["image_page"])
        mock_image_to_string = Mock(return_value="Picture Text")

        with patch(
            "codeqa.code_extractor.load_optional_ocr_modules",
            return_value=(mock_convert_from_path, mock_image_to_string),
        ):
            # Act
            text = extract_pdf_text(pdf_path)

        # Assert
        mock_convert_from_path.assert_called_once_with(str(pdf_path))
        mock_image_to_string.assert_called_once_with("image_page")
        self.assertIn("Picture Text", text)
